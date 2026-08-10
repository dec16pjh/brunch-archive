# -*- coding: utf-8 -*-
import os, json, datetime
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

base = os.path.dirname(__file__)
DATA_DIR = os.path.join(base, "full_data")
IMG_DIR = os.path.join(base, "images")
OUT_DIR = os.path.join(base, "output")
os.makedirs(OUT_DIR, exist_ok=True)

FONT = "Nanum Myeongjo"
AUTHOR = "박종호"


def set_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)


def fmt_date(ts):
    if not ts:
        return ""
    try:
        return datetime.datetime.fromtimestamp(ts / 1000).strftime("%Y.%m.%d")
    except Exception:
        return ""


def build_book(book):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(148)
    section.page_height = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    usable_w_mm = 148 - 18 - 18

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)

    # cover
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book["title"]); set_font(r, size=24, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(book.get("genre", "")); set_font(r, size=11, color=(0x66, 0x66, 0x66))
    for _ in range(10):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(AUTHOR); set_font(r, size=14, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("brunch.co.kr/@freeist"); set_font(r, size=9, color=(0x99, 0x99, 0x99))
    doc.add_page_break()

    # TOC
    p = doc.add_paragraph(); r = p.add_run("목차"); set_font(r, size=16, bold=True)
    doc.add_paragraph()
    current_part = None
    for ch in book["chapters"]:
        if ch["part"] and ch["part"] != current_part:
            current_part = ch["part"]
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
            r = p.add_run(current_part); set_font(r, size=11.5, bold=True, color=(0x33, 0x33, 0x33))
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Mm(6)
        r = p.add_run(f"{ch['seq']:02d}   {ch['title']}"); set_font(r, size=10.5)
    doc.add_page_break()

    # chapters
    current_part = None
    for ch in book["chapters"]:
        if ch["part"] and ch["part"] != current_part:
            if current_part is not None:
                doc.add_page_break()
            current_part = ch["part"]
            for _ in range(4):
                doc.add_paragraph()
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(current_part); set_font(r, size=18, bold=True)
            doc.add_page_break()
        else:
            doc.add_page_break()

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{ch['seq']:02d}"); set_font(r, size=10, color=(0xAA, 0xAA, 0xAA))
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(ch["title"]); set_font(r, size=17, bold=True)
        p.paragraph_format.space_after = Pt(4)
        if ch["subtitle"]:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(ch["subtitle"]); set_font(r, size=10.5, italic=True, color=(0x66, 0x66, 0x66))
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18)

        for block in ch["blocks"]:
            if block["type"] == "img":
                fname = block.get("local_file")
                if not fname:
                    continue
                img_path = os.path.join(IMG_DIR, fname)
                if not os.path.exists(img_path):
                    continue
                try:
                    from PIL import Image as PILImage
                    with PILImage.open(img_path) as im:
                        iw, ih = im.size
                    max_w_mm = usable_w_mm
                    max_h_mm = 150
                    w_mm = max_w_mm
                    h_mm = w_mm * ih / iw
                    if h_mm > max_h_mm:
                        h_mm = max_h_mm
                        w_mm = h_mm * iw / ih
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Mm(w_mm))
                    p.paragraph_format.space_after = Pt(10)
                except Exception:
                    continue
                continue

            text = block.get("text", "")
            if not text.strip():
                continue
            is_quote = block["type"] == "quote"
            lines = text.split("\n")
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.62
            p.paragraph_format.space_after = Pt(9)
            if is_quote:
                p.paragraph_format.left_indent = Mm(8)
                p.paragraph_format.right_indent = Mm(4)
            else:
                p.paragraph_format.first_line_indent = Mm(4.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for i, line in enumerate(lines):
                if i > 0:
                    p.add_run().add_break()
                r = p.add_run(line)
                set_font(r, size=10.3 if not is_quote else 9.8, italic=is_quote,
                         color=(0x55, 0x55, 0x55) if is_quote else None)

    footer = section.footer
    fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, "PAGE")
    header = section.header
    hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(book["title"]); set_font(r, size=8, color=(0xAA, 0xAA, 0xAA))

    out_path = os.path.join(OUT_DIR, f"{book['title']}.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    results = []
    for fn in sorted(os.listdir(DATA_DIR)):
        book = json.load(open(os.path.join(DATA_DIR, fn), encoding="utf-8"))
        out = build_book(book)
        results.append(out)
    with open(os.path.join(base, "docx_gen_log.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(results))
    print("DONE", len(results))
